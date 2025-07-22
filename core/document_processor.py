import os
import pandas as pd
import docx
import PyPDF2
from pathlib import Path
from typing import List, Dict, Any, Optional
from config.settings import settings
import logging

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Handle processing of various document types"""
    
    def __init__(self):
        self.supported_extensions = settings.ALLOWED_EXTENSIONS
    
    def process_document(self, file_path: str) -> Dict[str, Any]:
        """
        Process a document and extract text content
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Dictionary containing extracted content and metadata
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        extension = file_path.suffix.lower().lstrip('.')
        
        if extension not in self.supported_extensions:
            raise ValueError(f"Unsupported file type: {extension}")
        
        try:
            if extension == 'pdf':
                return self._process_pdf(file_path)
            elif extension in ['docx', 'doc']:
                return self._process_word(file_path)
            elif extension in ['xlsx', 'xls']:
                return self._process_excel(file_path)
            elif extension == 'csv':
                return self._process_csv(file_path)
            elif extension in ['txt', 'md']:
                return self._process_text(file_path)
            else:
                raise ValueError(f"Processing not implemented for: {extension}")
                
        except Exception as e:
            logger.error(f"Error processing {file_path}: {str(e)}")
            raise
    
    def _process_pdf(self, file_path: Path) -> Dict[str, Any]:
        """Extract text from PDF file"""
        text_content = []
        
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text.strip():
                        text_content.append(f"Page {page_num + 1}:\n{page_text}")
                except Exception as e:
                    logger.warning(f"Error extracting page {page_num + 1}: {str(e)}")
        
        return {
            'content': '\n\n'.join(text_content),
            'type': 'pdf',
            'pages': len(pdf_reader.pages),
            'filename': file_path.name,
            'size': file_path.stat().st_size
        }
    
    def _process_word(self, file_path: Path) -> Dict[str, Any]:
        """Extract text from Word document"""
        doc = docx.Document(file_path)
        
        # Extract paragraphs
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)
        
        # Extract tables
        tables_content = []
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            if table_data:
                tables_content.append(table_data)
        
        content = '\n\n'.join(paragraphs)
        if tables_content:
            content += "\n\nTables:\n" + str(tables_content)
        
        return {
            'content': content,
            'type': 'word',
            'paragraphs': len(paragraphs),
            'tables': len(tables_content),
            'filename': file_path.name,
            'size': file_path.stat().st_size
        }
    
    def _process_excel(self, file_path: Path) -> Dict[str, Any]:
        """Extract data from Excel file"""
        try:
            # Read all sheets
            excel_data = pd.read_excel(file_path, sheet_name=None, engine='openpyxl')
            
            content_parts = []
            summary = {
                'sheets': [],
                'total_rows': 0,
                'total_columns': 0
            }
            
            for sheet_name, df in excel_data.items():
                if not df.empty:
                    # Convert DataFrame to readable text
                    sheet_content = f"Sheet: {sheet_name}\n"
                    sheet_content += f"Columns: {', '.join(df.columns.astype(str))}\n"
                    sheet_content += f"Shape: {df.shape[0]} rows × {df.shape[1]} columns\n\n"
                    
                    # Add sample data (first 10 rows)
                    sheet_content += "Sample Data:\n"
                    sheet_content += df.head(10).to_string(index=False)
                    
                    # Add summary statistics for numeric columns
                    numeric_cols = df.select_dtypes(include=['number']).columns
                    if len(numeric_cols) > 0:
                        sheet_content += "\n\nNumeric Summary:\n"
                        sheet_content += df[numeric_cols].describe().to_string()
                    
                    content_parts.append(sheet_content)
                    
                    summary['sheets'].append({
                        'name': sheet_name,
                        'rows': df.shape[0],
                        'columns': df.shape[1],
                        'numeric_columns': len(numeric_cols)
                    })
                    summary['total_rows'] += df.shape[0]
                    summary['total_columns'] += df.shape[1]
            
            return {
                'content': '\n\n' + '='*50 + '\n\n'.join(content_parts),
                'type': 'excel',
                'summary': summary,
                'filename': file_path.name,
                'size': file_path.stat().st_size
            }
            
        except Exception as e:
            logger.error(f"Error processing Excel file: {str(e)}")
            raise
    
    def _process_csv(self, file_path: Path) -> Dict[str, Any]:
        """Extract data from CSV file"""
        try:
            df = pd.read_csv(file_path)
            
            content = f"CSV File: {file_path.name}\n"
            content += f"Columns: {', '.join(df.columns.astype(str))}\n"
            content += f"Shape: {df.shape[0]} rows × {df.shape[1]} columns\n\n"
            
            # Add sample data
            content += "Sample Data (first 10 rows):\n"
            content += df.head(10).to_string(index=False)
            
            # Add summary statistics
            numeric_cols = df.select_dtypes(include=['number']).columns
            if len(numeric_cols) > 0:
                content += "\n\nNumeric Summary:\n"
                content += df[numeric_cols].describe().to_string()
            
            return {
                'content': content,
                'type': 'csv',
                'rows': df.shape[0],
                'columns': df.shape[1],
                'numeric_columns': len(numeric_cols),
                'filename': file_path.name,
                'size': file_path.stat().st_size
            }
            
        except Exception as e:
            logger.error(f"Error processing CSV file: {str(e)}")
            raise
    
    def _process_text(self, file_path: Path) -> Dict[str, Any]:
        """Extract text from plain text file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
            
            return {
                'content': content,
                'type': 'text',
                'lines': len(content.splitlines()),
                'characters': len(content),
                'filename': file_path.name,
                'size': file_path.stat().st_size
            }
            
        except UnicodeDecodeError:
            # Try with different encoding
            with open(file_path, 'r', encoding='latin-1') as file:
                content = file.read()
                
            return {
                'content': content,
                'type': 'text',
                'lines': len(content.splitlines()),
                'characters': len(content),
                'filename': file_path.name,
                'size': file_path.stat().st_size,
                'encoding': 'latin-1'
            }
    
    def chunk_content(self, content: str, chunk_size: int = None, overlap: int = None) -> List[str]:
        """
        Split content into chunks for vector storage
        
        Args:
            content: Text content to chunk
            chunk_size: Size of each chunk
            overlap: Overlap between chunks
            
        Returns:
            List of text chunks
        """
        chunk_size = chunk_size or settings.CHUNK_SIZE
        overlap = overlap or settings.CHUNK_OVERLAP
        
        if len(content) <= chunk_size:
            return [content]
        
        chunks = []
        start = 0
        
        while start < len(content):
            end = start + chunk_size
            
            # Try to break at sentence boundaries
            if end < len(content):
                # Look for sentence endings within the last 100 characters
                search_start = max(end - 100, start)
                sentence_end = content.rfind('.', search_start, end)
                if sentence_end > start:
                    end = sentence_end + 1
            
            chunk = content[start:end].strip()
            if chunk:
                chunks.append(chunk)
            
            start = end - overlap if end < len(content) else end
        
        return chunks
